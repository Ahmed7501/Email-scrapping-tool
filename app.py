import streamlit as st
import tempfile
import os
import pandas as pd
from pathlib import Path
from main import EmailScraper
import time

def combine_results(all_results):
    """Combine results from multiple files."""
    combined = {
        'total_urls_processed': 0,
        'successful_scrapes': 0,
        'failed_scrapes': 0,
        'total_emails_found': 0,
        'unique_emails_found': 0,
        'success_rate': 0,
        'detailed_results': [],
        'output_files': {}
    }
    
    all_emails = set()
    
    for result in all_results:
        combined['total_urls_processed'] += result.get('total_urls_processed', 0)
        combined['successful_scrapes'] += result.get('successful_scrapes', 0)
        combined['failed_scrapes'] += result.get('failed_scrapes', 0)
        
        # Collect emails
        for detail_result in result.get('detailed_results', []):
            combined['detailed_results'].append(detail_result)
            for email in detail_result.get('emails', []):
                all_emails.add(email)
    
    combined['total_emails_found'] = len(all_emails)
    combined['unique_emails_found'] = len(all_emails)
    
    if combined['total_urls_processed'] > 0:
        combined['success_rate'] = (combined['successful_scrapes'] / combined['total_urls_processed']) * 100
    
    return combined

# Page configuration
st.set_page_config(
    page_title="📧 Email Scraper Tool",
    page_icon="📧",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 1rem;
    }
    .metric-card {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        border-left: 4px solid #1f77b4;
    }
    .success-message {
        background-color: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 0.5rem;
        border: 1px solid #c3e6cb;
    }
    .error-message {
        background-color: #f8d7da;
        color: #721c24;
        padding: 1rem;
        border-radius: 0.5rem;
        border: 1px solid #f5c6cb;
    }
</style>
""", unsafe_allow_html=True)

# Main header
st.markdown('<h1 class="main-header">📧 Email Scraper Tool</h1>', unsafe_allow_html=True)
st.markdown("""
<div style="text-align: center; margin-bottom: 2rem;">
    <p style="font-size: 1.1rem; color: #666;">
        Extract emails from websites and social media profiles. Upload files with URLs and get comprehensive results.
    </p>
</div>
""", unsafe_allow_html=True)

# Sidebar configuration
with st.sidebar:
    st.header("⚙️ Scraping Configuration")
    
    # Scraping options
    st.subheader("Options")
    use_selenium = st.checkbox("Use Selenium (for dynamic sites)", value=True, 
                              help="Enable for JavaScript-heavy websites")
    use_proxies = st.checkbox("Use Proxy Rotation", value=False,
                             help="Use proxy servers to avoid rate limiting")
    use_social = st.checkbox("Scrape Social Media", value=True,
                            help="Extract emails from social media profiles")
    max_internal = st.slider("Max Internal Pages", 1, 10, 3,
                            help="Number of internal pages to scrape per URL")
    output_format = st.selectbox("Output Format", ["csv", "excel", "both"], index=1)
    
    # Advanced options
    st.subheader("Advanced")
    timeout = st.slider("Request Timeout (seconds)", 10, 60, 30)
    delay = st.slider("Delay between requests (seconds)", 0.5, 3.0, 1.0, 0.1)
    
    st.markdown("---")
    st.markdown("**Tips:**")
    st.markdown("- Use Selenium for modern websites")
    st.markdown("- Enable proxies for large-scale scraping")
    st.markdown("- Higher delays reduce detection risk")

# Main content area
col1, col2 = st.columns([2, 1])

with col1:
    st.header("📁 Upload Files")
    
    # File uploader
    uploaded_files = st.file_uploader(
        "Upload files with URLs (.csv, .xlsx, .xls, .txt, .docx)",
        type=["csv", "xlsx", "xls", "txt", "docx"],
        accept_multiple_files=True,
        help="You can upload multiple files at once"
    )
    
    if uploaded_files:
        st.success(f"✅ Uploaded {len(uploaded_files)} file(s)")
        
        # Show file details
        file_details = []
        for file in uploaded_files:
            file_details.append({
                "Filename": file.name,
                "Size": f"{file.size / 1024:.1f} KB",
                "Type": file.type or "Unknown"
            })
        
        st.subheader("📋 File Details")
        st.dataframe(pd.DataFrame(file_details), use_container_width=True)

        # --- File Preview Section ---
        st.subheader("👀 File Preview")
        url_column_choices = {}
        for file in uploaded_files:
            st.markdown(f"**{file.name}**")
            try:
                file.seek(0)
                if file.name.endswith('.csv'):
                    df = pd.read_csv(file, nrows=10)
                    st.dataframe(df, use_container_width=True)
                    url_column_choices[file.name] = list(df.columns)
                elif file.name.endswith('.xlsx') or file.name.endswith('.xls'):
                    try:
                        df = pd.read_excel(file, engine='openpyxl', nrows=10)
                        st.dataframe(df, use_container_width=True)
                        url_column_choices[file.name] = list(df.columns)
                    except Exception as e:
                        st.error(f"Excel preview failed: {e}")
                elif file.name.endswith('.txt'):
                    content = file.read().decode('utf-8', errors='ignore').splitlines()
                    preview = '\n'.join(content[:10])
                    st.text(preview)
                elif file.name.endswith('.docx'):
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(file.read()))
                    text = '\n'.join([para.text for para in doc.paragraphs][:10])
                    st.text(text)
                else:
                    st.warning("Unsupported file type for preview.")
            except Exception as e:
                st.error(f"Preview failed: {e}")
        # Let user select URL column for each Excel/CSV file
        url_column_selection = {}
        for file in uploaded_files:
            if file.name.endswith('.csv') or file.name.endswith('.xlsx') or file.name.endswith('.xls'):
                columns = url_column_choices.get(file.name, [])
                if columns:
                    url_column_selection[file.name] = st.selectbox(f"Select URL column for {file.name}", columns, key=f"urlcol_{file.name}")

        # Before scraping, show extracted URLs for debug
        extracted_urls = []
        url_extraction_map = {}  # Map file name to list of URLs for scraping
        for file in uploaded_files:
            file.seek(0)
            try:
                if file.name.endswith('.csv'):
                    df = pd.read_csv(file)
                    col = url_column_selection.get(file.name, df.columns[0])
                    urls = df[col].dropna().astype(str).tolist()
                    extracted_urls.extend(urls)
                    url_extraction_map[file.name] = urls
                elif file.name.endswith('.xlsx') or file.name.endswith('.xls'):
                    df = pd.read_excel(file, engine='openpyxl')
                    col = url_column_selection.get(file.name, df.columns[0])
                    urls = df[col].dropna().astype(str).tolist()
                    extracted_urls.extend(urls)
                    url_extraction_map[file.name] = urls
                elif file.name.endswith('.txt'):
                    content = file.read().decode('utf-8', errors='ignore').splitlines()
                    urls = [line.strip() for line in content if line.strip()]
                    extracted_urls.extend(urls)
                    url_extraction_map[file.name] = urls
                elif file.name.endswith('.docx'):
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(file.read()))
                    urls = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                    extracted_urls.extend(urls)
                    url_extraction_map[file.name] = urls
            except Exception as e:
                st.error(f"Error extracting URLs from {file.name}: {e}")
        if extracted_urls:
            st.info(f"Extracted {len(extracted_urls)} URLs. Example: {extracted_urls[:5]}")
        else:
            st.warning("No URLs found in uploaded files.")

with col2:
    st.header("🚀 Start Scraping")
    
    if uploaded_files:
        if st.button("🔍 Start Email Scraping", type="primary", use_container_width=True):
            # Initialize session state
            if 'scraping_results' not in st.session_state:
                st.session_state.scraping_results = None
            if 'scraping_complete' not in st.session_state:
                st.session_state.scraping_complete = False
            
            progress_container = st.container()
            status_container = st.container()
            results_container = st.container()
            
            with progress_container:
                progress_bar = st.progress(0)
                status_text = st.empty()
            
            # Collect all URLs from url_extraction_map
            all_urls = []
            for urls in url_extraction_map.values():
                all_urls.extend(urls)
            
            if not all_urls:
                st.error("No URLs to scrape. Please upload files with valid URLs.")
            else:
                try:
                    with EmailScraper(
                        use_selenium=use_selenium,
                        use_proxies=use_proxies,
                        use_social_scraping=use_social,
                        max_internal_pages=max_internal,
                        output_format=output_format
                    ) as scraper:
                        # Scrape emails from the list of URLs
                        results = scraper.scrape_from_urls(all_urls)
                        progress_bar.progress(100)
                        status_text.success("Scraping complete!")
                        st.session_state.scraping_results = results
                        st.session_state.scraping_complete = True
                except Exception as e:
                    st.error(f"Error during scraping: {e}")
                    status_text.error("Scraping failed.")
                    st.session_state.scraping_complete = False

# Display results
if st.session_state.get('scraping_complete', False) and st.session_state.get('scraping_results'):
    results = st.session_state.scraping_results
    
    if results is not None:
        st.markdown("---")
        st.header("📊 Scraping Results")
        
        # Metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Total URLs", results.get('total_urls_processed', 0))
        
        with col2:
            st.metric("Successful", results.get('successful_scrapes', 0))
        
        with col3:
            st.metric("Unique Emails", results.get('unique_emails_found', 0))
        
        with col4:
            st.metric("Success Rate", f"{results.get('success_rate', 0):.1f}%")
        
        # Display emails in a table
        email_data = []
        for result in results.get('detailed_results', []):
            if result is not None:
                url = result.get('url', '')
                emails = result.get('emails', [])
                if emails:
                    for email in emails:
                        email_data.append({
                            'URL': url,
                            'Email': email,
                            'Source Page': result.get('source_page', ''),
                            'Status': result.get('status', ''),
                            'Source Type': result.get('source_type', '')
                        })
                else:
                    email_data.append({
                        'URL': url,
                        'Email': '',
                        'Source Page': result.get('source_page', ''),
                        'Status': result.get('status', ''),
                        'Source Type': result.get('source_type', '')
                    })
        if email_data:
            df = pd.DataFrame(email_data)
            st.dataframe(df, use_container_width=True)
        else:
            st.warning("No emails found for the provided URLs.")
            
            # Download options
            st.subheader("💾 Download Results")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # CSV download
                csv_data = df.to_csv(index=False)
                st.download_button(
                    label="📄 Download CSV",
                    data=csv_data,
                    file_name="extracted_emails.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            
            with col2:
                # Excel download
                try:
                    excel_buffer = pd.ExcelWriter('temp_emails.xlsx', engine='openpyxl')
                    df.to_excel(excel_buffer, index=False, sheet_name='Emails')
                    excel_buffer.close()
                    
                    with open('temp_emails.xlsx', 'rb') as f:
                        excel_data = f.read()
                    
                    st.download_button(
                        label="📊 Download Excel",
                        data=excel_data,
                        file_name="extracted_emails.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                    
                    # Clean up temp file
                    if os.path.exists('temp_emails.xlsx'):
                        os.unlink('temp_emails.xlsx')
                        
                except Exception as e:
                    st.error(f"Excel export failed: {e}")
            
            with col3:
                # Text download
                text_data = "\n".join([f"{row['Email']} - {row['Source URL']}" for _, row in df.iterrows()])
                st.download_button(
                    label="📝 Download Text",
                    data=text_data,
                    file_name="extracted_emails.txt",
                    mime="text/plain",
                    use_container_width=True
                )
    
    else:
        st.warning("⚠️ No emails found. Try enabling Selenium or checking your URLs.")
    
    # Show detailed results
    with st.expander("🔍 Detailed Results"):
        st.json(results)

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #666; padding: 1rem;">
    <p>Developed with ❤️ using Streamlit and Python</p>
    <p><a href="https://github.com/Ahmed7501/Email-scrapping-tool" target="_blank">View on GitHub</a></p>
</div>
""", unsafe_allow_html=True) 