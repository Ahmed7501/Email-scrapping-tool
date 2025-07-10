"""
File Parser Module

This module handles parsing of various file formats to extract URLs for scraping.
Supports CSV, Excel, TXT, and DOCX files.
"""

import csv
import logging
import re
from typing import List, Dict, Any
from pathlib import Path
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    print("Warning: pandas not available. CSV and TXT files only.")
from urllib.parse import urlparse, urljoin

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileParser:
    """
    A class for parsing various file formats and extracting URLs.
    
    Supports CSV, Excel (.xlsx, .xls), TXT, and DOCX files.
    """
    
    def __init__(self):
        """Initialize the FileParser with URL regex pattern."""
        # URL regex pattern for extraction
        self.url_pattern = re.compile(
            r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        )
        
        # Common URL column names in files
        self.url_column_names = {
            'url', 'website', 'link', 'site', 'domain', 'webpage',
            'address', 'web', 'page', 'homepage', 'home_page'
        }
    
    def parse_file(self, file_path: str) -> List[str]:
        """
        Parse a file and extract all URLs from it.
        
        Args:
            file_path (str): Path to the file to parse
            
        Returns:
            List[str]: List of unique URLs found in the file
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        path_obj = Path(file_path)
        
        if not path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = path_obj.suffix.lower()
        
        try:
            if file_extension == '.csv':
                return self._parse_csv(path_obj)
            elif file_extension in ['.xlsx', '.xls']:
                return self._parse_excel(path_obj)
            elif file_extension == '.txt':
                return self._parse_txt(path_obj)
            elif file_extension == '.docx':
                return self._parse_docx(path_obj)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            raise
    
    def _parse_csv(self, file_path: Path) -> List[str]:
        """Parse CSV file and extract URLs."""
        urls = []
        
        if PANDAS_AVAILABLE:
            try:
                # Try to read with pandas first (handles encoding better)
                df = pd.read_csv(file_path)
                urls.extend(self._extract_urls_from_dataframe(df))
                return self._clean_and_validate_urls(urls)
            except Exception as e:
                logger.warning(f"Pandas CSV parsing failed, trying with csv module: {e}")
        
        # Fallback to csv module
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                reader = csv.reader(file)
                for row in reader:
                    for cell in row:
                        urls.extend(self._extract_urls_from_text(cell))
        except Exception as e:
            logger.error(f"Error parsing CSV file: {e}")
        
        return self._clean_and_validate_urls(urls)
    
    def _parse_excel(self, file_path: Path) -> List[str]:
        """Parse Excel file and extract URLs."""
        urls = []
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                urls.extend(self._extract_urls_from_dataframe(df))
                
        except Exception as e:
            logger.error(f"Error parsing Excel file: {e}")
            raise
        
        return self._clean_and_validate_urls(urls)
    
    def _parse_txt(self, file_path: Path) -> List[str]:
        """Parse TXT file and extract URLs."""
        urls = []
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        urls.extend(self._extract_urls_from_text(content))
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise ValueError("Could not decode file with any supported encoding")
                
        except Exception as e:
            logger.error(f"Error parsing TXT file: {e}")
            raise
        
        return self._clean_and_validate_urls(urls)
    
    def _parse_docx(self, file_path: Path) -> List[str]:
        """Parse DOCX file and extract URLs."""
        urls = []
        
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                urls.extend(self._extract_urls_from_text(paragraph.text))
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        urls.extend(self._extract_urls_from_text(cell.text))
                        
        except ImportError:
            raise ValueError("python-docx library is required to parse DOCX files")
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {e}")
            raise
        
        return self._clean_and_validate_urls(urls)
    
    def _extract_urls_from_dataframe(self, df: pd.DataFrame) -> List[str]:
        """Extract URLs from a pandas DataFrame."""
        urls = []
        
        # Look for URL columns by name
        for col in df.columns:
            col_lower = col.lower().strip()
            if col_lower in self.url_column_names:
                # This column likely contains URLs
                for value in df[col].dropna():
                    urls.extend(self._extract_urls_from_text(str(value)))
        
        # If no URL columns found, search all columns
        if not urls:
            for col in df.columns:
                for value in df[col].dropna():
                    urls.extend(self._extract_urls_from_text(str(value)))
        
        return urls
    
    def _extract_urls_from_text(self, text: str) -> List[str]:
        """Extract URLs from text using regex."""
        if not text:
            return []
        
        matches = self.url_pattern.findall(text)
        return [url.strip() for url in matches if url.strip()]
    
    def _clean_and_validate_urls(self, urls: List[str]) -> List[str]:
        """
        Clean and validate URLs, removing duplicates and invalid ones.
        
        Args:
            urls (List[str]): List of raw URLs
            
        Returns:
            List[str]: List of clean, unique, valid URLs
        """
        clean_urls = []
        seen_urls = set()
        
        for url in urls:
            # Clean the URL
            url = url.strip()
            if not url:
                continue
            
            # Remove trailing punctuation
            url = url.rstrip('.,;:!?')
            
            # Ensure URL has protocol
            if not url.startswith(('http://', 'https://')):
                url = 'https://' + url
            
            # Validate URL format
            try:
                parsed = urlparse(url)
                if parsed.scheme and parsed.netloc:
                    # Normalize URL
                    normalized_url = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
                    if parsed.query:
                        normalized_url += f"?{parsed.query}"
                    
                    if normalized_url not in seen_urls:
                        clean_urls.append(normalized_url)
                        seen_urls.add(normalized_url)
            except Exception:
                logger.warning(f"Invalid URL format: {url}")
                continue
        
        logger.info(f"Extracted {len(clean_urls)} unique valid URLs")
        return clean_urls
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Get information about the file being parsed.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            Dict[str, Any]: File information including format, size, etc.
        """
        path_obj = Path(file_path)
        
        if not path_obj.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        info = {
            'name': path_obj.name,
            'extension': path_obj.suffix.lower(),
            'size_bytes': path_obj.stat().st_size,
            'size_mb': round(path_obj.stat().st_size / (1024 * 1024), 2),
            'exists': True
        }
        
        return info


# Convenience function for quick URL extraction
def extract_urls_from_file(file_path: str) -> List[str]:
    """
    Convenience function to extract URLs from a file.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        List[str]: List of extracted URLs
    """
    parser = FileParser()
    return parser.parse_file(file_path) 